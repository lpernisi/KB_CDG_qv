# -*- coding: utf-8 -*-
"""Convertitore Markdown -> Word (.docx) per le guide del progetto.
Gestisce il sottoinsieme di Markdown usato nei doc: titolo (#), sezioni (##/###),
sottotitolo in corsivo (*...*), elenchi puntati (con un livello di indentazione),
e formattazione inline **grassetto**, *corsivo*, `monospazio`.

Uso:  python tools/md_to_docx.py docs/guida-costo-prodotto.md [docs/guida-costo-prodotto.docx]
Poi (opzionale) l'export in PDF lo fa Word stesso (vedi tools/genera_guida.ps1).
"""
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches

INLINE = re.compile(r'(\*\*.+?\*\*|`[^`]+?`|\*[^*]+?\*)')

# Le emoji "semaforo" non sono nei font PDF standard (escono come quadratini):
# le rendiamo con un pallino "●" (presente in ogni font) colorato.
SEMAFORO = {
    '🟢': RGBColor(0x1A, 0x7F, 0x37),   # verde
    '🟡': RGBColor(0xC9, 0x8A, 0x00),   # ambra
    '🔴': RGBColor(0xC0, 0x39, 0x2B),   # rosso
}
EMOJI = re.compile('([%s])' % ''.join(SEMAFORO))


def _inline(p, text):
    for tok in INLINE.split(text):
        if not tok:
            continue
        if tok.startswith('**') and tok.endswith('**'):
            p.add_run(tok[2:-2]).bold = True
        elif tok.startswith('`') and tok.endswith('`'):
            r = p.add_run(tok[1:-1])
            r.font.name = 'Consolas'
            r.font.size = Pt(9.5)
            r.font.color.rgb = RGBColor(0x8A, 0x33, 0x24)
        elif tok.startswith('*') and tok.endswith('*'):
            p.add_run(tok[1:-1]).italic = True
        else:
            p.add_run(tok)


def add_runs(p, text):
    """Aggiunge i 'run' con la formattazione inline; le emoji semaforo -> pallino colorato."""
    for seg in EMOJI.split(text):
        if seg in SEMAFORO:
            r = p.add_run('●')        # ●
            r.bold = True
            r.font.color.rgb = SEMAFORO[seg]
        elif seg:
            _inline(p, seg)


def convert(src: Path, out: Path):
    md = src.read_text(encoding='utf-8')
    doc = Document()
    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(10.5)
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Inches(0.7)
        s.left_margin = s.right_margin = Inches(0.8)

    style_names = {s.name for s in doc.styles}
    para = []           # righe del paragrafo in costruzione
    state = {'list_p': None, 'quote': False}   # ultimo punto elenco (continuazioni) / paragrafo citazione

    def flush_para():
        if not para:
            return
        text = ' '.join(x.strip() for x in para).strip()
        para.clear()
        quote = state['quote']
        state['quote'] = False
        if not text:
            return
        # citazione/nota: tutto il paragrafo in corsivo, leggermente rientrato
        if quote:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            add_runs(p, text)
            for r in p.runs:
                r.italic = True
            return
        # sottotitolo: riga interamente in *corsivo* (singoli asterischi)
        if (text.startswith('*') and not text.startswith('**')
                and text.endswith('*') and not text.endswith('**')):
            p = doc.add_paragraph()
            if 'Subtitle' in style_names:
                p.style = doc.styles['Subtitle']
            add_runs(p, text[1:-1])
        else:
            add_runs(doc.add_paragraph(), text)

    for raw in md.split('\n'):
        line = raw.rstrip()
        s = line.strip()
        if s == '':
            flush_para()
            state['list_p'] = None
            continue
        for lvl_h, pref in ((0, '# '), (1, '## '), (2, '### ')):
            if s.startswith(pref):
                flush_para()
                state['list_p'] = None
                doc.add_heading(s[len(pref):].strip(), level=lvl_h)
                break
        else:
            mb = re.match(r'^(\s*)-\s+(.*)$', line)
            mo = re.match(r'^(\s*)\d+\.\s+(.*)$', line)
            if s.startswith('>'):
                flush_para()
                state['quote'] = True
                para.append(s.lstrip('>').strip())
            elif mb:
                flush_para()
                sn = 'List Bullet' if len(mb.group(1)) < 2 else 'List Bullet 2'
                p = doc.add_paragraph(style=sn if sn in style_names else None)
                add_runs(p, mb.group(2))
                state['list_p'] = p
            elif mo:
                flush_para()
                sn = 'List Number' if len(mo.group(1)) < 2 else 'List Number 2'
                p = doc.add_paragraph(style=sn if sn in style_names else None)
                add_runs(p, mo.group(2))
                state['list_p'] = p
            elif state['list_p'] is not None and not para:
                add_runs(state['list_p'], ' ' + s)   # continuazione del punto elenco
            else:
                para.append(s)
    flush_para()

    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    print(f'OK -> {out}')


if __name__ == '__main__':
    src = Path(sys.argv[1])
    out = Path(sys.argv[2]) if len(sys.argv) > 2 else src.with_suffix('.docx')
    convert(src, out)
